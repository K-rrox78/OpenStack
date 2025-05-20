import markdown
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os
import html

# Fonction pour ajouter un arrière-plan à un paragraphe
def add_shading_to_paragraph(paragraph, color):
    # Ajouter l'ombrage au paragraphe
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)  # Couleur d'arrière-plan
    paragraph._element.get_or_add_pPr().append(shading_elm)

# Chemin vers le fichier markdown
markdown_file = 'README.md'
output_docx = 'OpenStack_Authelia_Nginx_Installation_Final.docx'

# Lire le contenu du fichier markdown
with open(markdown_file, 'r', encoding='utf-8') as file:
    md_content = file.read()

# Extraire les titres et créer une structure de sommaire
headings = []
for line in md_content.split('\n'):
    if line.startswith('## '):
        headings.append(('h2', line[3:].strip()))
    elif line.startswith('### '):
        headings.append(('h3', line[4:].strip()))

# Créer un nouveau document Word
doc = Document()

# Configurer les marges du document
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Titre principal
title = None
for line in md_content.split('\n'):
    if line.startswith('# '):
        title = line[2:]
        break

if title:
    title_heading = doc.add_heading(title, level=0)
    title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Ajouter une page de garde
doc.add_paragraph()
doc.add_paragraph("Document d'Architecture Technique", style='Subtitle')
doc.add_paragraph("Date: Mai 2025", style='Subtitle')
doc.add_paragraph()

# Ajouter un saut de page
doc.add_page_break()

# Titre du sommaire
toc_title = doc.add_heading("Sommaire", level=1)
toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# Générer manuellement la table des matières
for level, title in headings:
    p = doc.add_paragraph()
    if level == 'h2':
        p.style = 'TOC 1'
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
    elif level == 'h3':
        p.style = 'TOC 2'
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(title)
        run.font.size = Pt(11)
        
    # Ajouter une ligne de pointillés (facultatif)
    for i in range(50):  # Nombre de points approximatif
        p.add_run('.')
        
    # Ajouter un numéro de page factice (pour la mise en page)
    p.add_run(' X')

# Ajouter un saut de page après le sommaire
doc.add_page_break()

# Analyser le contenu Markdown
lines = md_content.split('\n')
i = 0

while i < len(lines):
    line = lines[i].strip()
    
    # Traitement des titres
    if line.startswith('## '):
        # Titre de niveau 2
        heading_text = line[3:]
        heading = doc.add_heading(heading_text, level=1)
        heading.style.font.size = Pt(16)
        heading.style.font.color.rgb = RGBColor(0, 0, 128)  # Bleu foncé
    
    elif line.startswith('### '):
        # Titre de niveau 3
        heading_text = line[4:]
        heading = doc.add_heading(heading_text, level=2)
        heading.style.font.size = Pt(14)
        heading.style.font.color.rgb = RGBColor(0, 0, 100)  # Bleu foncé
    
    # Traitement des listes
    elif line.startswith('- '):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(line[2:])
    
    # Traitement des blocs de code
    elif line.startswith('```'):
        code_language = line[3:].strip()
        code_block = []
        i += 1
        
        # Collecter toutes les lignes du bloc de code
        while i < len(lines) and not lines[i].startswith('```'):
            code_block.append(lines[i])
            i += 1
        
        # Ajouter le bloc de code au document avec fond coloré
        if code_block:
            p = doc.add_paragraph()
            code_text = '\n'.join(code_block)
            code_run = p.add_run(code_text)
            code_run.font.name = 'Courier New'
            code_run.font.size = Pt(9)
            code_run.font.color.rgb = RGBColor(0, 0, 0)  # Texte noir
            
            # Ajouter l'ombrage au paragraphe (fond gris clair)
            add_shading_to_paragraph(p, "EEEEEE")  # Code hexadécimal pour gris clair
            
            # Ajouter une ligne vide après le bloc de code
            doc.add_paragraph()
    
    # Traitement des paragraphes normaux
    elif line and not line.startswith('#') and not line.startswith('-'):
        # Vérifier s'il s'agit d'un paragraphe multi-lignes
        paragraph_text = [line]
        next_i = i + 1
        while next_i < len(lines) and lines[next_i].strip() and not lines[next_i].startswith('#') and not lines[next_i].startswith('-') and not lines[next_i].startswith('```'):
            paragraph_text.append(lines[next_i].strip())
            next_i += 1
        
        # Ajouter le paragraphe
        if paragraph_text:
            p = doc.add_paragraph()
            p.add_run(' '.join(paragraph_text))
            
            # Avancer i pour éviter de traiter à nouveau ces lignes
            i = next_i - 1
    
    i += 1

# Enregistrer le document
doc.save(output_docx)

print(f"Le document avec sommaire manuel a été créé: '{output_docx}'")
