import markdown
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

# Chemin vers le fichier markdown
markdown_file = 'README.md'
output_docx = 'OpenStack_Authelia_Nginx_Installation.docx'

# Lire le contenu du fichier markdown
with open(markdown_file, 'r', encoding='utf-8') as file:
    md_content = file.read()

# Convertir le markdown en HTML
html = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])

# Créer un nouveau document Word
doc = Document()

# Configurer les marges du document
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Analyser le HTML pour le convertir en document Word
# Extraire les différentes parties
title_pattern = re.compile(r'<h1>(.*?)</h1>')
h2_pattern = re.compile(r'<h2>(.*?)</h2>')
h3_pattern = re.compile(r'<h3>(.*?)</h3>')
p_pattern = re.compile(r'<p>(.*?)</p>', re.DOTALL)
ul_pattern = re.compile(r'<ul>(.*?)</ul>', re.DOTALL)
li_pattern = re.compile(r'<li>(.*?)</li>')
code_pattern = re.compile(r'<pre><code>(.*?)</code></pre>', re.DOTALL)

# Extraire le titre principal
title_matches = title_pattern.findall(html)
if title_matches:
    title = doc.add_heading(title_matches[0], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Parcourir le HTML et le convertir en éléments Word
lines = html.split('\n')
i = 0
while i < len(lines):
    line = lines[i]
    
    # Traitement des titres h2
    h2_match = h2_pattern.search(line)
    if h2_match:
        heading = doc.add_heading(h2_match.group(1), level=1)
        heading.style.font.size = Pt(16)
        heading.style.font.color.rgb = RGBColor(0, 0, 128)  # Bleu foncé
        i += 1
        continue
    
    # Traitement des titres h3
    h3_match = h3_pattern.search(line)
    if h3_match:
        heading = doc.add_heading(h3_match.group(1), level=2)
        heading.style.font.size = Pt(14)
        heading.style.font.color.rgb = RGBColor(0, 0, 100)  # Bleu foncé
        i += 1
        continue
    
    # Traitement des paragraphes
    p_match = p_pattern.search(line)
    if p_match:
        p = doc.add_paragraph()
        p.add_run(p_match.group(1))
        i += 1
        continue
    
    # Traitement des listes
    ul_match = ul_pattern.search(line)
    if ul_match:
        ul_content = ul_match.group(1)
        li_items = li_pattern.findall(ul_content)
        for item in li_items:
            p = doc.add_paragraph(item, style='List Bullet')
        i += 1
        continue
    
    # Traitement des blocs de code
    code_match = code_pattern.search(line)
    if code_match and i < len(lines):
        code_content = code_match.group(1)
        # Remplacer les entités HTML pour les caractères spéciaux
        code_content = code_content.replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
        p = doc.add_paragraph()
        code_run = p.add_run(code_content)
        code_run.font.name = 'Courier New'
        code_run.font.size = Pt(9)
        
        # Ajouter une ligne vide après le bloc de code
        doc.add_paragraph()
        i += 1
        continue
    
    i += 1

# Enregistrer le document
doc.save(output_docx)

print(f"Le document a été converti avec succès et enregistré sous '{output_docx}'")
