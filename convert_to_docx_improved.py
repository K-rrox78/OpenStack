import markdown
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os
import html

# Chemin vers le fichier markdown
markdown_file = 'README.md'
output_docx = 'OpenStack_Authelia_Nginx_Installation_Complete.docx'

# Lire le contenu du fichier markdown
with open(markdown_file, 'r', encoding='utf-8') as file:
    md_content = file.read()

# Convertir le markdown en HTML
html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code', 'codehilite'])

# Créer un nouveau document Word
doc = Document()

# Configurer les marges du document
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Fonction pour extraire du texte HTML 
def clean_html_text(text):
    text = re.sub(r'<.*?>', '', text)
    return html.unescape(text)

# Fonction pour traiter les sections de code
def process_code_block(code_text):
    # Nettoyer le texte du code
    code_text = html.unescape(code_text)
    lines = code_text.split('\n')
    
    # Retirer le premier/dernier élément s'ils sont vides
    if lines and not lines[0].strip():
        lines = lines[1:]
    if lines and not lines[-1].strip():
        lines = lines[:-1]
    
    return '\n'.join(lines)

# Analyser le contenu Markdown
lines = md_content.split('\n')
i = 0

# Titre principal
title = None
for line in lines:
    if line.startswith('# '):
        title = line[2:]
        break

if title:
    title_heading = doc.add_heading(title, level=0)
    title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Table des matières
toc_heading = doc.add_heading("Table des matières", level=1)
toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()  # Espace après le titre

i = 0
while i < len(lines):
    line = lines[i].strip()
    
    # Traitement des titres
    if line.startswith('## '):
        # Titre de niveau 2
        heading_text = line[3:]
        heading = doc.add_heading(heading_text, level=1)
        heading.style.font.size = Pt(16)
        heading.style.font.color.rgb = RGBColor(0, 0, 128)  # Bleu foncé
    
    elif line.startswith('### '):
        # Titre de niveau 3
        heading_text = line[4:]
        heading = doc.add_heading(heading_text, level=2)
        heading.style.font.size = Pt(14)
        heading.style.font.color.rgb = RGBColor(0, 0, 100)  # Bleu foncé
    
    # Traitement des listes
    elif line.startswith('- '):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(line[2:])
    
    # Traitement des blocs de code
    elif line.startswith('```'):
        code_language = line[3:].strip()
        code_block = []
        i += 1
        
        # Collecter toutes les lignes du bloc de code
        while i < len(lines) and not lines[i].startswith('```'):
            code_block.append(lines[i])
            i += 1
        
        # Ajouter le bloc de code au document
        if code_block:
            p = doc.add_paragraph()
            code_text = '\n'.join(code_block)
            code_run = p.add_run(code_text)
            code_run.font.name = 'Courier New'
            code_run.font.size = Pt(9)
            
            # Ajouter une ligne vide après le bloc de code
            doc.add_paragraph()
    
    # Traitement des paragraphes normaux
    elif line and not line.startswith('#') and not line.startswith('-'):
        # Vérifier s'il s'agit d'un paragraphe multi-lignes
        paragraph_text = [line]
        next_i = i + 1
        while next_i < len(lines) and lines[next_i].strip() and not lines[next_i].startswith('#') and not lines[next_i].startswith('-') and not lines[next_i].startswith('```'):
            paragraph_text.append(lines[next_i].strip())
            next_i += 1
        
        # Ajouter le paragraphe
        if paragraph_text:
            p = doc.add_paragraph()
            p.add_run(' '.join(paragraph_text))
            
            # Avancer i pour éviter de traiter à nouveau ces lignes
            i = next_i - 1
    
    i += 1

# Enregistrer le document
doc.save(output_docx)

print(f"Le document a été converti avec succès et enregistré sous '{output_docx}'")
